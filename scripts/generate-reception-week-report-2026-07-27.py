#!/usr/bin/env python3
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
LOGO = ROOT / "wp-content/themes/eduma-child/assets/images/toolkit-logo.png"
PLAN = ROOT / "reports/weekly-plans/Toolkit_Weekly_Work_Plan_27_31_July_2026.docx"
MILESTONE = ROOT / "reports/weekly-milestones/Toolkit_Weekly_Milestone_13_Reception_Integration_and_Deployment_Readiness_31_July_2026.docx"
OLIVE, ORANGE, TEAL, DARK, LIGHT = "969E2A", "FF6600", "006A68", "262B2A", "F4F6F2"


def rgb(value):
    return RGBColor.from_string(value)


def shade(cell, value):
    props = cell._tc.get_or_add_tcPr()
    item = OxmlElement("w:shd")
    item.set(qn("w:fill"), value)
    props.append(item)


def setup(title, header_text, footer_text):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Inches(0.62)
    section.left_margin = section.right_margin = Inches(0.7)
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.2)
    normal.font.color.rgb = rgb(DARK)
    for name, size, color in (("Title", 25, TEAL), ("Heading 1", 16, TEAL), ("Heading 2", 12, OLIVE)):
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
    header = section.header.paragraphs[0]
    header.text = header_text
    header.runs[0].bold = True
    header.runs[0].font.color.rgb = rgb(TEAL)
    header.runs[0].font.size = Pt(8)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(footer_text).font.size = Pt(7.5)
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.add_run().add_picture(str(LOGO), width=Inches(1.5))
    heading = doc.add_paragraph(style="Title")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.add_run(title)
    return doc


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def table(doc, headers, rows, color=TEAL, widths=None):
    result = doc.add_table(rows=1, cols=len(headers))
    result.style = "Table Grid"
    result.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, label in zip(result.rows[0].cells, headers):
        shade(cell, color)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = cell.paragraphs[0].add_run(label)
        run.bold = True
        run.font.color.rgb = rgb("FFFFFF")
    for values in rows:
        cells = result.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    return result


def build_plan():
    doc = setup(
        "Weekly Work Plan",
        "TOOLKIT AFRICA  /  WEEKLY WORK PLAN  /  27-31 JULY 2026",
        "Toolkit Africa | Reception Integration, Demo Acceptance and Production Readiness | Internal working document",
    )
    subtitle = doc.add_paragraph("27-31 JULY 2026")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.color.rgb = rgb(ORANGE)
    focus = doc.add_paragraph("Reception System  |  Website Form  |  Secure Data Relay  |  Demo  |  Production")
    focus.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Purpose", level=1)
    doc.add_paragraph(
        "Move the locally verified website-to-reception workflow through a controlled hosting, demo and production-readiness process. "
        "The week prioritizes visitor-data protection, operational usability, recoverable deployment and evidence-based acceptance."
    )

    doc.add_heading("Verified starting position", level=1)
    bullets(doc, [
        "A branded WordPress reception form is implemented at the database-independent /reception/ route.",
        "WordPress validates and rate-limits submissions, then relays them to Laravel using a timestamped HMAC signature and one-use nonce.",
        "WordPress retains no second copy of visitor personal data; the reception application remains authoritative.",
        "Website records enter Reception Control as source=website and status=follow_up with no physical check-in timestamp.",
        "Reception metrics exclude website requests until a staff member records an actual arrival.",
        "An automated local WordPress-to-Laravel request passed and labelled test data was removed; no local browser-review server remains running.",
        "The Laravel suite passes 18 tests and 92 assertions; Pint, PHP syntax, JavaScript syntax and whitespace checks pass.",
        "Independent demo/production roots now exist, but the partially deployed demo returns HTTP 500 until its private environment, database and runtime setup are completed.",
    ])

    doc.add_heading("Daily plan", level=1)
    table(doc, ("Day", "Primary work", "Required output", "Release gate"), [
        ("Mon 27", "Consolidate the integration contract, local evidence, field mapping, privacy boundary, migration and rollback procedure.", "Approved technical report, changed-file inventory, passing test record and local database backup.", "No remote change; no secrets or PII in Git, reports or browser configuration."),
        ("Tue 28", "Confirm demo and production reception hostnames, private application paths, public/ document roots, databases, HTTPS, backups and cron capability.", "Hosting topology record, access validation and least-privilege deployment checklist.", "Do not deploy Laravel into a WordPress public directory or accept invalid TLS."),
        ("Wed 29", "Resume the partially deployed reception backend: configure a demo-only environment/database, resolve the current HTTP 500, migrate, and verify staff authentication and signed API controls.", "Healthy HTTP 200 demo backend, migration evidence, replay rejection, staff register visibility and rollback archive.", "Backend health, HTTPS, private-file protection and database backup must pass."),
        ("Thu 30", "Deploy the reviewed WordPress child-theme files to demo with the form initially disabled; configure the private API URL/secret, enable last, purge cache and test desktop/mobile.", "Accepted demo /reception/ page, successful labelled request, staff follow-up record, accessibility and failure-path evidence.", "No main-domain work before demo acceptance; remove labelled test records after inspection."),
        ("Fri 31", "Prepare or execute the main release only after approval. Repeat backups, use a new production secret, deploy backend first, enable the form last and monitor.", "Production readiness decision or verified production release, smoke-test record, monitoring snapshot and WMR-13 milestone close.", "Named authorization, rollback owner, same-day backups, valid TLS and demo sign-off required."),
    ])

    doc.add_heading("Acceptance criteria", level=1)
    table(doc, ("Area", "Completion test", "Control"), [
        ("Public form", "Required fields, optional fields, consent, mobile layout, keyboard use and clear success/error states pass.", "No sensitive data in URLs, analytics, browser logs or WordPress storage."),
        ("Signed relay", "Valid requests return a WEB reference; stale, changed-signature and replayed requests are rejected.", "Different 32+ character secrets for demo and production; server-side storage only."),
        ("Reception Control", "Website request is visible, filterable and editable by authorized staff.", "Source remains identifiable and staff access continues to be audited."),
        ("Attendance integrity", "Website follow-ups do not affect check-in totals, QR share or purpose metrics before physical arrival.", "Staff transition sets check-in time only when checked-in/checked-out status is selected."),
        ("Operations", "Database backup, migration, scheduler, cache, logs and rollback are verified.", "No PII request bodies in logs; encrypted fields and approved retention remain active."),
        ("Deployment", "Demo passes before main; backend is healthy before the website form is enabled.", "Fail closed with the reception phone number if the backend or configuration is unavailable."),
    ], OLIVE)

    doc.add_heading("Required owner and hosting inputs", level=1)
    bullets(doc, [
        "Demo and production reception hostnames.",
        "Private Laravel application directories with document roots mapped only to public/.",
        "Demo and production database credentials, backup method and restore authority.",
        "Certificate-valid HTTPS and a scheduler/cron facility.",
        "Named demo checker, production authorizer and rollback owner.",
        "Privacy/retention confirmation and authorized reception staff accounts.",
    ])

    doc.add_heading("Immediate rollback", level=1)
    doc.add_paragraph(
        "Set TOOLKIT_RECEPTION_FORM_ENABLED to false and purge LiteSpeed. The public page will retain the phone fallback and stop sending data. "
        "Restore the previous child-theme package only for a presentation defect. Roll back Laravel independently; do not reverse the nullable "
        "check-in migration while website follow-up records exist. Rotate the shared secret if exposure is suspected."
    )
    PLAN.parent.mkdir(parents=True, exist_ok=True)
    doc.save(PLAN)


def build_milestone():
    doc = setup(
        "Reception Integration and Deployment Readiness",
        "TOOLKIT AFRICA  /  WEEKLY MILESTONE 13  /  27-31 JULY 2026",
        "Toolkit Africa | WMR-13 | Reception Integration and Deployment Readiness",
    )
    subtitle = doc.add_paragraph("WMR-13  |  REPORTING PERIOD: 27-31 JULY 2026")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.color.rgb = rgb(ORANGE)
    status = doc.add_paragraph("STATUS AT OPENING: LOCAL IMPLEMENTATION COMPLETE; REMOTE RELEASE GATED")
    status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    status.runs[0].bold = True

    doc.add_heading("Executive milestone", level=1)
    doc.add_paragraph(
        "Establish a secure public reception-request journey that feeds Toolkit Africa's existing Reception Control system without duplicating "
        "visitor data in WordPress, while preserving the distinction between advance follow-up and physical attendance. The milestone closes "
        "when the workflow is accepted on demo and is either safely promoted to production or formally held with every unmet gate recorded."
    )

    doc.add_heading("Document control", level=1)
    table(doc, ("Field", "Value"), [
        ("Report", "Toolkit Weekly Milestone Report 13"),
        ("Reporting period", "27-31 July 2026"),
        ("Prepared", "27 July 2026"),
        ("Workstream", "Reception system and public website integration"),
        ("Systems", "WordPress Eduma child theme and Laravel Reception Control"),
        ("Data classification", "Internal report; contains no visitor PII or deployment secrets"),
    ])

    doc.add_heading("Milestone deliverables", level=1)
    table(doc, ("Deliverable", "Opening status", "Week-close evidence"), [
        ("Website reception experience", "Implemented locally", "Responsive /reception/ page, accessibility review and accepted visitor guidance."),
        ("Secure server relay", "Implemented and locally verified", "Signed demo request, replay rejection, rate-limit and failure-path results."),
        ("Reception data model", "Migration implemented locally", "Backed-up demo migration and verified source/status/timestamp behavior."),
        ("Staff operations", "Views and metrics updated locally", "Authorized staff can review/update requests without attendance distortion."),
        ("Demo release", "Pending hosting topology", "Backend and WordPress release manifests, smoke tests, test-record cleanup and sign-off."),
        ("Production release", "Not started; demo-gated", "Go/no-go decision, approvals, backups, smoke tests, monitoring and rollback evidence."),
        ("Operational handover", "Drafted", "Configuration, deployment, monitoring, retention and rollback instructions accepted."),
    ], OLIVE)

    doc.add_heading("Completed foundation at milestone opening", level=1)
    bullets(doc, [
        "Added a signed Laravel endpoint for website-originated reception requests.",
        "Added timestamp tolerance, nonce replay protection and HMAC verification over the exact raw JSON body.",
        "Added nullable check-in and website submission fields through a reversible database migration.",
        "Added the WordPress same-origin form relay with nonce checking, honeypot protection and per-IP throttling.",
        "Added a branded public reception form using the reception system's defined visitor fields.",
        "Updated Reception Control lists, details, dashboard totals, exports and retention handling.",
        "Verified a complete local submission and removed the labelled test record afterward.",
        "Recorded configuration, deployment, acceptance and rollback instructions without secrets.",
    ])

    doc.add_heading("Validation baseline", level=1)
    bullets(doc, [
        "Laravel: 18 tests passed with 92 assertions.",
        "Laravel Pint: passed.",
        "PHP syntax: reception integration, theme functions and reception template passed.",
        "JavaScript syntax: reception form controller passed.",
        "Automated local WordPress pretty-permalink check returned HTTP 200; no persistent local review environment was provided.",
        "Automated local WordPress REST relay returned HTTP 200 with a generated WEB reference.",
        "Stored record: source website, status follow_up, submitted timestamp present and checked-in timestamp absent.",
        "Labelled integration records: deleted after verification.",
    ])

    doc.add_heading("Risks and controls for the reporting week", level=1)
    table(doc, ("Risk", "Impact", "Control / decision"), [
        ("Laravel deployed beneath a public WordPress root", "Sensitive configuration or application files could be exposed.", "Require a private app path and a document root mapped exactly to public/."),
        ("Shared secret exposed or reused", "Unauthorized request creation across environments.", "Use independent random secrets, store server-side only, rotate on suspicion."),
        ("Website leads counted as attendance", "Incorrect operational reporting.", "Keep checked_in_at null until staff records physical arrival; exclude website source from attendance metrics."),
        ("Duplicate PII storage", "Expanded privacy and breach surface.", "Reception remains authoritative; WordPress performs no fallback persistence."),
        ("Partial deployment", "Public form fails or loses submissions.", "Deploy backend first, test health, enable WordPress form last, fail closed to phone."),
        ("Unrecoverable migration", "Service interruption or data loss.", "Same-day database backup, migration evidence and tested rollback procedure."),
    ])

    doc.add_heading("Current deployment checkpoint", level=1)
    bullets(doc, [
        "Independent cPanel application/public directories were created for demo and production outside WordPress.",
        "reception-demo.toolkitafrica.ac.ke was created and mapped to the independent demo public directory.",
        "Laravel commit e8bf6cc was uploaded and extracted; temporary installers, archives and cron entries were removed.",
        "The demo currently returns HTTP 500 and is not accepted because private environment, database, migrations and runtime permissions remain incomplete.",
        "The reception form and redesigned image/video gallery templates are local-only and have not been deployed to demo.",
        "No persistent local browser-review server or locally reviewable form exists.",
        "Production remains unchanged apart from empty independent directory creation.",
    ])

    doc.add_heading("Week-close status fields", level=1)
    doc.add_paragraph(
        "Complete these fields on 31 July: demo release result; demo acceptance owner; production go/no-go decision; production authorization; "
        "backup identifiers; deployed release identifiers; smoke-test result; labelled test-record cleanup; monitoring outcome; rollback readiness; "
        "open risks and next dated action."
    )

    doc.add_heading("Definition of done", level=1)
    bullets(doc, [
        "Demo and production hosting boundaries are documented and private files are confirmed inaccessible.",
        "Demo backend, form, staff workflow, metrics, security controls and rollback pass.",
        "Production is changed only after explicit authorization and same-day backups.",
        "The production form is enabled only after backend health and signed submission checks pass.",
        "Reception staff can identify and action website requests without confusing them with physical visits.",
        "No secrets, visitor PII or request bodies are present in Git, milestone reports or analytics.",
        "If any gate remains unmet, the milestone records a hold rather than claiming deployment completion.",
    ])

    MILESTONE.parent.mkdir(parents=True, exist_ok=True)
    doc.save(MILESTONE)


build_plan()
build_milestone()
print(PLAN)
print(MILESTONE)

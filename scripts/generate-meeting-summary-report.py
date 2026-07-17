#!/usr/bin/env python3
"""Generate the concise Toolkit Africa website-modernisation meeting report."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "reports" / "meeting-summary"
ASSETS = OUT / "assets"
EVIDENCE = ROOT / "reports" / "weekly-milestones" / "evidence"
LOGO = ROOT / "wp-content" / "themes" / "eduma-child" / "assets" / "images" / "toolkit-logo.png"

OLIVE = "969E2A"
ORANGE = "FF6600"
TEAL = "006A68"
DARK = "262B2A"
LIGHT = "F4F6F2"
LINE = "D9DED6"


def rgb(value):
    return RGBColor.from_string(value)


def shade(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    item = props.find(qn("w:shd")) or OxmlElement("w:shd")
    item.set(qn("w:fill"), fill)
    if item.getparent() is None:
        props.append(item)


def border(cell, color=LINE):
    props = cell._tc.get_or_add_tcPr()
    borders = props.first_child_found_in("w:tcBorders") or OxmlElement("w:tcBorders")
    if borders.getparent() is None:
        props.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        item = OxmlElement("w:" + edge)
        item.set(qn("w:val"), "single")
        item.set(qn("w:sz"), "6")
        item.set(qn("w:color"), color)
        borders.append(item)


def optimize(source, target, max_width=1400):
    with Image.open(source) as image:
        image = image.convert("RGB")
        if image.width > max_width:
            height = round(image.height * max_width / image.width)
            image = image.resize((max_width, height), Image.Resampling.LANCZOS)
        image.save(target, "JPEG", quality=78, optimize=True, progressive=True)


def configure(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)
    normal.font.color.rgb = rgb(DARK)
    normal.paragraph_format.space_after = Pt(6)
    for name, size, color in (("Title", 27, TEAL), ("Heading 1", 20, TEAL), ("Heading 2", 13, OLIVE)):
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(6)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("TOOLKIT AFRICA  /  MEETING BRIEF")
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = rgb(ORANGE)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Website Modernisation Programme | 17 July 2026 | Internal meeting document")
    run.font.size = Pt(7.5)
    run.font.color.rgb = rgb(TEAL)


def title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(LOGO), width=Inches(1.65))
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("EXECUTIVE MEETING SUMMARY")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = rgb(ORANGE)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Website Modernisation\nProgramme Update")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Milestones from June to 17 July 2026")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = rgb(OLIVE)
    box = doc.add_table(rows=1, cols=1)
    shade(box.cell(0, 0), LIGHT)
    border(box.cell(0, 0), OLIVE)
    p = box.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Purpose\n").bold = True
    p.add_run("Summarise the new demo experience, delivery progress, corrected course-page coverage, SEO and measurement foundations, and the controlled route to main-domain rollout.")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Prepared for project and management review").italic = True
    doc.add_page_break()


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.22)
        p.add_run(item)


def comparison(doc, name, label):
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, (heading, path, fill) in enumerate((
        ("CURRENT MAIN SITE", EVIDENCE / f"current-{name}.jpg", TEAL),
        ("NEW DEMO SITE", EVIDENCE / f"demo-{name}.jpg", OLIVE),
    )):
        cell = table.cell(0, index)
        shade(cell, fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(heading)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = rgb("FFFFFF")
        body = table.cell(1, index)
        body.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        body.paragraphs[0].add_run().add_picture(str(path), width=Inches(3.2))
        body.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        border(cell)
        border(body)
    caption = doc.add_paragraph(label)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].italic = True
    caption.runs[0].font.size = Pt(8)


def status_table(doc):
    rows = (
        ("Homepage", "Modernised", "Hero, 2026 banner, Who We Are, impact metrics, video, pathways, testimonials and assistant."),
        ("About Us", "Modernised", "Institutional story, programme context and all 18 current team members."),
        ("Courses", "Corrected", "Directory plus eight public legacy course URLs now use the shared modern detail structure."),
        ("Blog / Notice Board", "Modernised", "Styled content discovery, working search and category filters."),
        ("Apply / Contact", "Modernised", "Clear admissions route, support details and measured conversion interactions."),
        ("SEO / Analytics", "Implemented", "Curated metadata, schema support, canonical handling and first-party engagement metrics."),
    )
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for idx, label in enumerate(("Area", "Status", "Meeting note")):
        cell = table.rows[0].cells[idx]
        shade(cell, TEAL)
        run = cell.paragraphs[0].add_run(label)
        run.bold = True
        run.font.color.rgb = rgb("FFFFFF")
    for area, status, note in rows:
        cells = table.add_row().cells
        cells[0].paragraphs[0].add_run(area).bold = True
        cells[1].paragraphs[0].add_run(status).font.color.rgb = rgb(OLIVE)
        cells[2].paragraphs[0].add_run(note)


def asset_page(doc):
    doc.add_heading("Outreach Materials and Visual Assets", level=1)
    doc.add_paragraph("Existing campaign material supports the wider communications ecosystem. It is referenced here as a supporting asset source, not as a replacement for the website's approved content hierarchy.")
    table = doc.add_table(rows=2, cols=3)
    items = (
        (ASSETS / "poster-tvet.jpg", "TVET/VTC poster", "Training outreach"),
        (ASSETS / "poster-welding.jpg", "University welding poster", "Welding awareness"),
        (ASSETS / "poster-opportunity.jpg", "Opportunity poster", "Campaign support"),
    )
    for index, (path, title, note) in enumerate(items):
        top = table.cell(0, index)
        top.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        top.paragraphs[0].add_run().add_picture(str(path), width=Inches(2.0))
        bottom = table.cell(1, index)
        p = bottom.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title + "\n")
        run.bold = True
        run.font.color.rgb = rgb(OLIVE)
        p.add_run(note)
        border(top)
        border(bottom)
    doc.add_heading("Additional approved-use candidates", level=2)
    bullets(doc, [
        "High-resolution solar training photography for renewable-energy course pages and admissions stories.",
        "Virtual-reality workshop photography for advanced welding and VR-supported learning pages.",
        "Graduation photography for learner outcomes, impact stories and testimonial context.",
        "All selected web assets should be cropped for context, resized responsively, compressed, and supplied with accurate alternative text before publication.",
    ])
    doc.add_page_break()


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    ASSETS.mkdir(parents=True, exist_ok=True)
    sources = {
        "poster-tvet.jpg": Path("/home/t316/Desktop/extracts/TVET-VTC.png"),
        "poster-welding.jpg": Path("/home/t316/Desktop/extracts/Universtiy_welding.png"),
        "poster-opportunity.jpg": Path("/home/t316/Desktop/Projects_father/toolkit/imgs/posters/pref1.jpeg"),
        "solar-training.jpg": Path("/home/t316/Desktop/extracts/Solar/Toolkit 003.jpg"),
        "vr-training.jpg": Path("/home/t316/Desktop/extracts/virtual reality/20260127_122403.jpg"),
        "graduation.jpg": Path("/home/t316/Desktop/extracts/GRADUATION/KMM_2143.JPG"),
    }
    for name, source in sources.items():
        optimize(source, ASSETS / name)

    doc = Document()
    configure(doc)
    title_page(doc)

    doc.add_heading("Executive Position", level=1)
    doc.add_paragraph("The demo site has moved from a builder-heavy, visually inconsistent experience to a child-theme implementation with controlled templates, a defined brand system and measurable user journeys. The main domain has not been changed.")
    status_table(doc)
    doc.add_heading("Immediate correction completed", level=2)
    doc.add_paragraph("All eight public course-related URLs now resolve through the modern course-page renderer locally. This closes the gap where only Welding had received the intended redesign while other pages remained blank or legacy-styled.")
    doc.add_page_break()

    doc.add_heading("Experience Transformation", level=1)
    comparison(doc, "home", "Homepage comparison: clearer hierarchy, deeper engagement flow and brand-led calls to action.")
    doc.add_paragraph()
    comparison(doc, "about", "About comparison: stronger institutional narrative and a complete, modern Our Team presentation.")
    doc.add_page_break()

    doc.add_heading("Courses and Admissions", level=1)
    comparison(doc, "courses", "Course discovery comparison: focused directory, accurate imagery and consistent programme navigation.")
    doc.add_heading("Course pages covered", level=2)
    bullets(doc, [
        "Welding and Fabrication; Renewable Energy; Organic Farming Skills; Digital Skills.",
        "Recognition of Prior Learning; Consultancy and Research; Online Training Portal - Jielimishe.",
        "Training Welders with Virtual Reality, with welding-specific imagery and learning focus.",
        "Course names and existing URLs are retained. Pricing remains controlled separately so unapproved future fees are not exposed.",
    ])
    doc.add_page_break()

    doc.add_heading("SEO, Performance and Measurement", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(0, 0).paragraphs[0].add_run().add_picture(str(EVIDENCE / "seo-courses.jpg"), width=Inches(3.25))
    p = table.cell(0, 1).paragraphs[0]
    p.add_run("Implemented foundation\n").bold = True
    p.runs[0].font.color.rgb = rgb(TEAL)
    p.add_run("Curated page titles and descriptions, canonical course URLs, breadcrumb hierarchy, social imagery and schema integration support stronger search interpretation and sharing.")
    bullets_in_cell = (
        "Legacy builder assets removed from rebuilt page payloads.",
        "First-party page view, scroll depth, dwell, interaction and performance events.",
        "Toolkit Control panel for feature switches and aggregate site metrics.",
    )
    for item in bullets_in_cell:
        q = table.cell(0, 1).add_paragraph(style="List Bullet")
        q.add_run(item)
    doc.add_heading("Course-page verification - 17 July 2026", level=2)
    doc.add_paragraph("Eight of eight local course URLs returned HTTP 200 with modern markup. Warm responses measured approximately 0.35-0.47 seconds and page HTML reduced to about 54 KB from roughly 155-205 KB on the legacy render. These are local smoke-test figures, not production Core Web Vitals.")
    doc.add_page_break()

    asset_page(doc)

    doc.add_heading("Rollout Controls and Meeting Decisions", level=1)
    doc.add_heading("Deployment safeguards", level=2)
    bullets(doc, [
        "Demo changes are deployed only after local route, asset, syntax and visual checks.",
        "A persistent latest-demo rollback is retained locally outside temporary storage.",
        "The redesign, 2026 catalogue and future pricing each have separate switches.",
        "The main-domain redesign defaults off and can be returned to the existing site immediately without deleting WordPress content.",
        "Main-domain rollout requires a fresh files/database backup, route acceptance checks and explicit authorization.",
    ])
    doc.add_heading("Items requiring management confirmation", level=2)
    bullets(doc, [
        "Approve the corrected demo course-page presentation after visual review.",
        "Confirm final course visibility and current admissions wording before main-domain cutover.",
        "Approve the Mzizi field contract and a non-production submission test before building the application adapter.",
        "Schedule the controlled main-domain release only after demo acceptance is complete.",
    ])
    doc.add_paragraph()
    box = doc.add_table(rows=1, cols=1)
    shade(box.cell(0, 0), OLIVE)
    p = box.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Recommended meeting outcome: approve the demo correction for inspection, retain the pricing gate, and continue through the documented main-domain readiness checklist.")
    run.bold = True
    run.font.color.rgb = rgb("FFFFFF")

    target = OUT / "Toolkit_Website_Modernisation_Meeting_Summary_17_July_2026.docx"
    doc.save(target)
    print(target)


if __name__ == "__main__":
    main()

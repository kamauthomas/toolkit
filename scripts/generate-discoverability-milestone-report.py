#!/usr/bin/env python3
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "reports/weekly-milestones/Toolkit_Weekly_Milestone_09_Search_Discoverability_and_Entity_SEO_17_July_2026.docx"
LOGO = ROOT / "wp-content/themes/eduma-child/assets/images/toolkit-logo.png"
OLIVE, ORANGE, TEAL, DARK, LIGHT = "969E2A", "FF6600", "006A68", "262B2A", "F4F6F2"


def rgb(value):
    return RGBColor.from_string(value)


def shade(cell, value):
    props = cell._tc.get_or_add_tcPr()
    item = OxmlElement("w:shd")
    item.set(qn("w:fill"), value)
    props.append(item)


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


doc = Document()
section = doc.sections[0]
section.top_margin = section.bottom_margin = Inches(0.65)
section.left_margin = section.right_margin = Inches(0.72)
normal = doc.styles["Normal"]
normal.font.name = "Aptos"
normal.font.size = Pt(9.5)
normal.font.color.rgb = rgb(DARK)
for name, size, color in (("Title", 26, TEAL), ("Heading 1", 17, TEAL), ("Heading 2", 12.5, OLIVE)):
    style = doc.styles[name]
    style.font.name = "Aptos Display"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = rgb(color)

header = section.header.paragraphs[0]
header.text = "TOOLKIT AFRICA  /  WEBSITE MODERNISATION  /  WMR-09"
header.runs[0].bold = True
header.runs[0].font.color.rgb = rgb(TEAL)
header.runs[0].font.size = Pt(8)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("Toolkit Africa | Search Discoverability and Entity SEO | 17 July 2026 | Confidential project record").font.size = Pt(7.5)

cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.add_run().add_picture(str(LOGO), width=Inches(1.55))
title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("Search Discoverability\nand Entity SEO")
subtitle = doc.add_paragraph("WEEKLY MILESTONE REPORT 09  |  17 JULY 2026")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].bold = True
subtitle.runs[0].font.color.rgb = rgb(ORANGE)

doc.add_heading("Executive summary", level=1)
doc.add_paragraph(
    "This milestone strengthens the relationship between the Toolkit name, Toolkit Africa, its official digital profiles, and its practical training programmes. The implementation prioritises branded and course-intent searches such as Toolkit Africa, Toolkit welding, MIG/MAG welding training, and MIG/MAG welding training in Kenya. Search position cannot be guaranteed for the globally ambiguous single word 'toolkit'; the work establishes the technical and content signals required to compete for relevant branded intent."
)

doc.add_heading("Implementation completed", level=1)
bullets(doc, [
    "Added Organization and EducationalOrganization structured data with Toolkit Africa's official name, recognised aliases, logo, contact details, Kikuyu location, and social profiles.",
    "Added WebSite entity data identifying Toolkit and Toolkit for Skills and Innovation as alternate names of the Toolkit Africa website.",
    "Added Course structured data to enabled modern and legacy course pages, connecting each course to Toolkit Africa as its provider.",
    "Optimised the MIG/MAG course title and meta description for practical training intent, Kenya relevance, Toolkit branding, and the Kikuyu location.",
    "Removed remaining public MIG/MAG wording that described the programme as fabrication and aligned headings and outcomes with the approved course name.",
    "Added visible course FAQs covering location, current fees and intakes, and application steps without publishing unconfirmed pricing.",
    "Expanded llms.txt discovery content with Toolkit brand aliases and a direct canonical reference to MIG/MAG Welding training.",
])

doc.add_heading("Search target map", level=1)
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for cell, label in zip(table.rows[0].cells, ("Search intent", "Primary destination", "Signal implemented")):
    shade(cell, TEAL)
    run = cell.paragraphs[0].add_run(label)
    run.bold = True
    run.font.color.rgb = rgb("FFFFFF")
rows = [
    ("Toolkit / Toolkit Africa", "Homepage", "Organization, WebSite aliases, official profiles"),
    ("Toolkit courses", "Course directory", "Internal hierarchy and provider relationship"),
    ("Toolkit welding", "MIG/MAG Welding page", "Dedicated metadata, H1, Course schema"),
    ("MIG MAG welding Kenya", "MIG/MAG Welding page", "Training intent, Kenya and Kikuyu relevance"),
    ("Toolkit course fees / intake", "Course pages and Admissions", "Visible FAQ and admissions verification"),
]
for values in rows:
    cells = table.add_row().cells
    for cell, value in zip(cells, values):
        cell.text = value

doc.add_heading("Validation record", level=1)
bullets(doc, [
    "All modified PHP files passed php -l syntax validation.",
    "The rendered MIG/MAG page returned HTTP 200 in the local WordPress environment.",
    "Rendered title verified: MIG/MAG Welding Training in Kenya | Toolkit Africa.",
    "Rendered description verified with practical MIG/MAG training, Toolkit Africa, Kikuyu, Kenya, workshop and VR-supported learning terms.",
    "Rendered JSON-LD verified for Organization, EducationalOrganization, official aliases and Course provider relationships.",
    "Git whitespace validation completed with no theme-source errors.",
])

doc.add_heading("Indexing and authority follow-up", level=1)
doc.add_paragraph(
    "Technical SEO makes the pages understandable and eligible for indexing; ranking for broad terms also depends on Google recrawling the main domain, consistent organisation profiles, Search Console ownership, local business signals, authoritative references and relevant backlinks. After main-domain rollout, the XML sitemap should be resubmitted in Google Search Console and priority URLs should be requested for indexing."
)

doc.add_heading("Governance", level=1)
bullets(doc, [
    "Course names remain controlled by the approved Toolkit catalogue.",
    "Fees, schedules and availability remain subject to Admissions confirmation.",
    "No future pricing was exposed through metadata, visible FAQs or structured data.",
    "Existing public URLs were retained to avoid unnecessary redirects and loss of accumulated search signals.",
])

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)

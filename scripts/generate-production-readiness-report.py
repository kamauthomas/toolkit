#!/usr/bin/env python3
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "reports/weekly-milestones/Toolkit_Weekly_Milestone_10_Production_Readiness_Audit_20_July_2026.docx"
LOGO = ROOT / "wp-content/themes/eduma-child/assets/images/toolkit-logo.png"
SHOTS = Path("/tmp/toolkit-audit-shots")
OLIVE, ORANGE, TEAL, DARK, PALE = "969E2A", "FF6600", "006A68", "252A29", "F4F6F2"


def rgb(value):
    return RGBColor.from_string(value)


def shade(cell, value):
    props = cell._tc.get_or_add_tcPr()
    item = OxmlElement("w:shd")
    item.set(qn("w:fill"), value)
    props.append(item)


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


doc = Document()
section = doc.sections[0]
section.top_margin = section.bottom_margin = Inches(0.62)
section.left_margin = section.right_margin = Inches(0.72)
normal = doc.styles["Normal"]
normal.font.name = "Aptos"
normal.font.size = Pt(9.3)
normal.font.color.rgb = rgb(DARK)
for name, size, color in (("Title", 24, TEAL), ("Heading 1", 16, TEAL), ("Heading 2", 12, OLIVE)):
    style = doc.styles[name]
    style.font.name = "Aptos Display"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = rgb(color)

header = section.header.paragraphs[0]
header.text = "TOOLKIT AFRICA  /  WEEKLY MILESTONE 10  /  20 JULY 2026"
header.runs[0].bold = True
header.runs[0].font.size = Pt(8)
header.runs[0].font.color.rgb = rgb(TEAL)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("Toolkit Africa | Production Readiness Audit | Internal release record").font.size = Pt(7.5)

cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.add_run().add_picture(str(LOGO), width=Inches(1.45))
title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("Production Readiness Audit")
subtitle = doc.add_paragraph("WEEKLY MILESTONE 10  |  20 JULY 2026")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].bold = True
subtitle.runs[0].font.color.rgb = rgb(ORANGE)
decision = doc.add_paragraph("DECISION: CONDITIONAL NO-GO UNTIL PRODUCTION GATES ARE COMPLETE")
decision.alignment = WD_ALIGN_PARAGRAPH.CENTER
decision.runs[0].bold = True
decision.runs[0].font.color.rgb = rgb(OLIVE)

doc.add_heading("Milestone objective", level=1)
doc.add_paragraph(
    "Complete the final consistency, security, responsive-layout and rollout audit of the demo release; correct verified defects; preserve a working rollback; and define the exact controls required before the main-domain cutover."
)

doc.add_heading("Work completed", level=1)
rows = [
    ("Public journeys", "Homepage, Notice Board, application route and all eight preserved course URLs return expected content."),
    ("Retired pages", "Students Portal now redirects permanently to Our Courses and is excluded from sitemap discovery. Duplicate Courses and Blog routes also redirect to their canonical destinations."),
    ("Responsive header", "Critical mobile drawer rules prevent the menu panel from appearing over content while parent-theme CSS is deferred."),
    ("Media", "The homepage story control now loads the approved YouTube no-cookie player on demand and removes it when closed."),
    ("Data accuracy", "Unsupported September-intake and language claims were removed from the Notice Board; current course names and preserved URLs remain intact."),
    ("Security", "Anonymous REST user listing and XML-RPC are blocked. Security headers were expanded. Metrics now require same-host origin/referrer and restrict path input."),
    ("SEO privacy", "Retired utility pages are noindexed/excluded and operational staff author profiles are removed from sitemap discovery."),
    ("Rollback", "Every changed demo child-theme file was backed up to persistent `rollbacks/latest-demo` before upload."),
]
table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for cell, label in zip(table.rows[0].cells, ("Area", "Verified milestone result")):
    shade(cell, TEAL)
    run = cell.paragraphs[0].add_run(label)
    run.bold = True
    run.font.color.rgb = rgb("FFFFFF")
for area, result in rows:
    cells = table.add_row().cells
    cells[0].text = area
    cells[1].text = result

doc.add_heading("Verification results", level=1)
bullets(doc, [
    "Demo homepage: HTTP 200; five configured browser security headers present.",
    "Students Portal and legacy Courses routes: HTTP 301 to /our-ventures/.",
    "Anonymous REST users: HTTP 404. Valid XML-RPC method request: HTTP 403.",
    "Originless metrics event: HTTP 403. Same-origin metrics event: HTTP 204.",
    "Retired page entries in Yoast page sitemap: zero. Staff author sitemap entries: excluded.",
    "Eight of eight current course routes: HTTP 200, distinct H1, populated image markup.",
    "Modified PHP syntax checks and Git whitespace/error validation: pass.",
])

home_shot = SHOTS / "demo-home-mobile-final.png"
notice_shot = SHOTS / "demo-notice-tablet-final.png"
if home_shot.exists() and notice_shot.exists():
    doc.add_heading("Responsive evidence", level=1)
    doc.add_paragraph("Deployed demo captures after the patch. The mobile header remains compact and the Notice Board controls retain a usable tablet layout.")
    evidence = doc.add_table(rows=1, cols=2)
    evidence.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, path, caption in zip(evidence.rows[0].cells, (home_shot, notice_shot), ("Homepage - mobile", "Notice Board - tablet")):
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(path), width=Inches(2.75))
        text = cell.add_paragraph(caption)
        text.alignment = WD_ALIGN_PARAGRAPH.CENTER
        text.runs[0].bold = True

doc.add_heading("Remaining production gates", level=1)
bullets(doc, [
    "Take and verify same-day production database and full-files backups outside a temporary directory.",
    "Verify main-host transfer access and TLS. The demo certificate exception must not be repeated for production.",
    "Inventory production wp-config.php, active theme, plugins, uploads, LiteSpeed rules, permalinks and rollout-switch values.",
    "Synchronize only the reviewed child theme with redesign, 2026 catalogue and September pricing switches off; do not reinstall WordPress or import the demo database.",
    "Confirm the legacy presentation after sync, enable only the redesign switch, purge cache and complete the production smoke test.",
    "Test instant rollback by disabling the redesign switch and purging cache before closing the maintenance window.",
    "Resolve whether Research, Toolkit in Brief, TTI Media and Gallery are modernized, consolidated by tested 301, or retained and improved before indexing.",
    "Verify production canonicals, sitemap, robots, llms files, contact mail, Mzizi handoff, chatbot, metrics, mobile/desktop layouts and cache hits after activation.",
])

doc.add_heading("WordPress preservation decision", level=1)
doc.add_paragraph(
    "The main-domain rollout does not require a WordPress reinstall. The existing core, database, uploads, plugins and parent theme remain in place. The child-theme release is added while switches are off, then the presentation is activated through TOOLKIT_REDESIGN_ENABLED. Disabling that switch and purging cache restores the previous presentation without moving or deleting site data."
)

doc.add_heading("Release recommendation", level=1)
doc.add_paragraph(
    "The demo patch is suitable for stakeholder review. Production remains a conditional no-go until the backup, access, legacy-state, content-route and rollback gates are evidenced and approved."
)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)

#!/usr/bin/env python3
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "reports/weekly-plans/Toolkit_Weekly_Work_Plan_20_24_July_2026.docx"
LOGO = ROOT / "wp-content/themes/eduma-child/assets/images/toolkit-logo.png"
OLIVE, ORANGE, TEAL, DARK, LIGHT = "969E2A", "FF6600", "006A68", "262B2A", "F4F6F2"


def rgb(value):
    return RGBColor.from_string(value)


def shade(cell, value):
    props = cell._tc.get_or_add_tcPr()
    item = OxmlElement("w:shd")
    item.set(qn("w:fill"), value)
    props.append(item)


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


doc = Document()
section = doc.sections[0]
section.top_margin = section.bottom_margin = Inches(0.62)
section.left_margin = section.right_margin = Inches(0.7)
normal = doc.styles["Normal"]
normal.font.name = "Aptos"
normal.font.size = Pt(9.2)
normal.font.color.rgb = rgb(DARK)
for name, size, color in (("Title", 25, TEAL), ("Heading 1", 16, TEAL), ("Heading 2", 12, OLIVE)):
    style = doc.styles[name]
    style.font.name = "Aptos Display"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = rgb(color)

header = section.header.paragraphs[0]
header.text = "TOOLKIT AFRICA  /  WEEKLY WORK PLAN  /  20-24 JULY 2026"
header.runs[0].bold = True
header.runs[0].font.color.rgb = rgb(TEAL)
header.runs[0].font.size = Pt(8)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("Toolkit Africa | Production Readiness, Outreach and Discoverability | Internal working document").font.size = Pt(7.5)

cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.add_run().add_picture(str(LOGO), width=Inches(1.5))
title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("Weekly Work Plan")
subtitle = doc.add_paragraph("20-24 JULY 2026")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].bold = True
subtitle.runs[0].font.color.rgb = rgb(ORANGE)
focus = doc.add_paragraph("Bulk SMS  |  Posters  |  SEO  |  Content  |  Production Cutover")
focus.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading("Purpose", level=1)
doc.add_paragraph(
    "Move the accepted demo release toward a controlled main-domain rollout while preparing factual outreach content. Work is gated by data approval, recipient consent, complete backups, demo acceptance, and a tested instant rollback."
)

doc.add_heading("Verified baseline", level=1)
bullets(doc, [
    "The redesign is active on demo.toolkitafrica.ac.ke; the main domain has not been replaced by this release.",
    "Latest recorded mobile audit: performance 84, FCP 1.7 s, LCP 3.1 s, TBT 60 ms, and CLS 0.045.",
    "Organization/Course schema, branded metadata, course FAQs, AI discovery files, and first-party interaction metrics are implemented.",
    "The existing WordPress database, Eduma parent theme, uploads, plugins, posts and Elementor content remain available.",
    "No SMS provider, sender ID, consented recipient list, delivery account, or approved campaign schedule is documented yet.",
    "Poster and training-photo source assets exist, but publication approval is not established for every file.",
])

doc.add_heading("Daily plan", level=1)
table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ("Day", "Primary work", "Required output", "Release gate")
for cell, label in zip(table.rows[0].cells, headers):
    shade(cell, TEAL)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    run = cell.paragraphs[0].add_run(label)
    run.bold = True
    run.font.color.rgb = rgb("FFFFFF")
rows = [
    ("Mon 20", "Confirm courses, fees, intakes and campaign audience. Inventory posters. Audit SMS provider, sender, consent and opt-out. Crawl demo.", "Approved data sheet, asset inventory, SMS requirements and release manifest.", "No unverified price, date, accreditation or outcome claim."),
    ("Tue 21", "Draft SMS variants. Design approved web/social/print poster formats. Select course-accurate photos and optimize web derivatives.", "Campaign matrix, message pack, approved creative set and asset manifest.", "Owner approval and accurate destination links."),
    ("Wed 22", "Apply factual content updates. Validate metadata, headings, schema, canonicals, sitemap, AI discovery, responsive images and performance.", "Release candidate and complete demo acceptance record.", "All priority pages, forms, assets, mobile and desktop pass."),
    ("Thu 23", "Take production backups. Sync child theme with all switches off. Test legacy state. Enable redesign only, purge cache and smoke test.", "Release record, backups, changed-file manifest and rollback proof.", "Explicit production authorization and rollback owner available."),
    ("Fri 24", "Monitor errors and engagement. Verify cache hits. Submit sitemap/index requests. Run internal SMS test. Publish only approved posters.", "Monitoring snapshot, SMS test result, indexing record and fresh milestone report.", "No broad SMS send until internal test and compliance checks pass."),
]
for values in rows:
    cells = table.add_row().cells
    for cell, value in zip(cells, values):
        cell.text = value
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

doc.add_heading("Production preservation model", level=1)
doc.add_paragraph(
    "The rollout will not reinstall WordPress. It will not import the demo database, replace WordPress core, change DNS/site URLs, or overwrite plugins and uploads. The approved child theme is synchronized while the redesign, 2026 catalog and pricing switches remain off. After legacy-state testing, only TOOLKIT_REDESIGN_ENABLED is activated. The original database content and parent-theme presentation remain available."
)
doc.add_paragraph(
    "Immediate presentation rollback: set TOOLKIT_REDESIGN_ENABLED to false and purge LiteSpeed. File rollback: restore the latest pre-cutover child-theme archive. Database restoration is reserved for a verified database-caused incident."
)

doc.add_heading("Workstream controls", level=1)
controls = doc.add_table(rows=1, cols=3)
controls.style = "Table Grid"
controls.alignment = WD_TABLE_ALIGNMENT.CENTER
for cell, label in zip(controls.rows[0].cells, ("Workstream", "This week's completion test", "Control")):
    shade(cell, OLIVE)
    run = cell.paragraphs[0].add_run(label)
    run.bold = True
    run.font.color.rgb = rgb("FFFFFF")
for values in [
    ("Bulk SMS", "Approved copy and successful internal test with delivery evidence.", "Consent, sender identity, opt-out, least data, no recipient data in Git or reports."),
    ("Posters", "Approved dimensions, text, imagery, links and optimized exports.", "Course-accurate assets and approved brand colours; retain print master separately."),
    ("SEO", "Priority routes have correct metadata, schema, canonical, sitemap and indexing handoff.", "No guaranteed-ranking claims; measure verified impressions, clicks and positions."),
    ("Content", "Owners approve changed facts and pages pass responsive review.", "Preserve valid existing copy; no future fees before the approved date."),
    ("Production", "New UI passes smoke tests and the old UI can be restored by switch.", "Same-day backups, child-theme-only sync, switches off during upload, documented rollback."),
]:
    cells = controls.add_row().cells
    for cell, value in zip(cells, values):
        cell.text = value

doc.add_heading("Production go/no-go checklist", level=1)
bullets(doc, [
    "Written acceptance of the demo release and explicit authorization for the main-domain cutover.",
    "Correct production access and TLS/FTPS certificate verification; do not repeat the demo certificate exception.",
    "Same-day full files and database backups stored persistently and tested for readability.",
    "Recorded current production state, release commit, file manifest and rollback owner.",
    "Legacy site passes after file sync with all three switches false.",
    "Homepage, About, courses, every active course route, Blog, Notice Board, application, Contact, privacy, login/account and 404 pass after activation.",
    "Logo/media, canonical hostname, forms, tracking, chatbot, console, mobile layout and second-request cache hits verified.",
])

doc.add_heading("Items requiring owner input", level=1)
bullets(doc, [
    "Approved course/intake information and whether any prices may appear this week.",
    "Bulk-SMS provider, sender ID, lawful/consented list owner, opt-out method and campaign authorization.",
    "Poster campaign priority, channels, final copy owner and image approvals.",
    "Production hosting credentials, verified backup mechanism, maintenance window and rollback decision-maker.",
    "Google Search Console and analytics access for verified indexing and performance records.",
])

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)

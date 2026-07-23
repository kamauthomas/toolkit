#!/usr/bin/env python3
"""Generate dated Toolkit milestone reports without merging reporting periods."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "reports" / "weekly-milestones"
LOGO = ROOT / "wp-content" / "themes" / "eduma-child" / "assets" / "images" / "toolkit-logo.png"

OLIVE = "969E2A"
ORANGE = "FF6600"
TEAL = "006A68"
DARK = "25281F"
MUTED = "606657"
LIGHT = "F4F6EF"


def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    fill = OxmlElement("w:shd")
    fill.set(qn("w:fill"), color)
    tc_pr.append(fill)


def set_cell_text(cell, text, bold=False, color=DARK, size=9):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_header(doc, report_id, date):
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(7.0))
    table.columns[0].width = Inches(1.25)
    table.columns[1].width = Inches(5.75)
    table.cell(0, 0).paragraphs[0].add_run().add_picture(str(LOGO), width=Inches(0.72))
    p = table.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"{report_id}\n{date}")
    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(TEAL)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Toolkit Africa | Website Modernisation Programme | Internal Milestone Record").font.size = Pt(8)


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(4)
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string(DARK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.name = "Aptos"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(OLIVE)


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor.from_string(TEAL)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.08
    if bold_prefix and text.startswith(bold_prefix):
        lead, rest = text.split(":", 1)
        run = p.add_run(lead + ":")
        run.bold = True
        p.add_run(rest)
    else:
        p.add_run(text)
    for run in p.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor.from_string(DARK)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        run.font.name = "Aptos"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor.from_string(DARK)


def add_control_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    set_cell_text(table.cell(0, 0), "Control", True, "FFFFFF")
    set_cell_text(table.cell(0, 1), "Recorded state", True, "FFFFFF")
    shade(table.cell(0, 0), TEAL)
    shade(table.cell(0, 1), TEAL)
    for index, (label, value) in enumerate(rows):
        cells = table.add_row().cells
        set_cell_text(cells[0], label, True)
        set_cell_text(cells[1], value)
        if index % 2 == 0:
            shade(cells[0], LIGHT)
            shade(cells[1], LIGHT)
    return table


def add_evidence(doc, path, caption):
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(6.55))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.italic = True
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def document(report_id, date, title, subtitle):
    doc = Document()
    add_header(doc, report_id, date)
    add_title(doc, title, subtitle)
    return doc


def production_report():
    doc = document(
        "WMR-11",
        "20 July 2026",
        "Production Activation and Main-Domain Stabilisation",
        "Controlled child-theme rollout, rollback preservation, cache correction, and responsive navigation acceptance",
    )
    add_heading(doc, "Executive milestone")
    add_body(doc, "The approved modern child theme was activated on the existing main-domain WordPress installation. WordPress core, the production database, plugins, uploads, parent theme, and legacy content were retained. The rollout switch remains the immediate presentation rollback mechanism.")
    add_heading(doc, "Completed work")
    add_bullets(doc, [
        "Synchronized the reviewed Eduma child theme through certificate-verified FTPS without reinstalling WordPress.",
        "Enabled the modern redesign while retaining the 2026 catalogue and September pricing switches in the disabled state.",
        "Purged a stale LiteSpeed object that caused the unparameterized public URL to continue serving the legacy interface.",
        "Corrected desktop navigation flow, duplicate dropdown presentation, sticky-header spacing, logo separation, right-side balance, and mobile logo clipping.",
        "Preserved production rollback copies and exact desktop/mobile screenshot evidence outside temporary storage.",
        "Confirmed main-domain modern pages, course routes, logo, security headers, REST user protection, assistant surface, and HTTPS assets.",
    ])
    add_heading(doc, "Deployment controls")
    add_control_table(doc, [
        ("WordPress installation", "Retained; no reinstall"),
        ("Production database", "Retained; demo database not imported"),
        ("Modern redesign", "Enabled"),
        ("2026 catalogue", "Disabled"),
        ("September pricing", "Disabled"),
        ("Immediate rollback", "Disable redesign switch and purge LiteSpeed"),
        ("Accepted commit", "7f5c5b3 - Finalize production header layout"),
    ])
    add_heading(doc, "Validation and evidence")
    add_bullets(doc, [
        "Exact public homepage verified after cache purge, including subsequent LiteSpeed cache hits.",
        "Desktop acceptance covered 1440 and 1366 pixel layouts; mobile acceptance covered 390 pixels.",
        "HSTS, nosniff, frame, referrer, and permissions headers remained present.",
        "Temporary cPanel purge jobs, markers, activation scripts, and test credentials were removed after use.",
    ])
    add_evidence(
        doc,
        ROOT / "rollbacks" / "production-20260720" / "screenshots" / "home-header-balanced-desktop.png",
        "Accepted main-domain desktop header and hero, captured 20 July 2026.",
    )
    add_evidence(
        doc,
        ROOT / "rollbacks" / "production-20260720" / "screenshots" / "home-header-balanced-mobile.png",
        "Accepted main-domain mobile header and hero, captured 20 July 2026.",
    )
    add_heading(doc, "Residual work at close of date")
    add_bullets(doc, [
        "Authenticated visual acceptance of future Toolkit Control changes requires an existing WordPress administrator session.",
        "Search Console and Bing Webmaster ownership, query reporting, and search-performance ingestion remain separate follow-up work.",
        "The application adapter was not part of this production activation.",
    ])
    doc.save(OUTPUT / "Toolkit_Weekly_Milestone_11_Production_Activation_20_July_2026.docx")


def support_report():
    doc = document(
        "WMR-12",
        "23 July 2026",
        "Enquiries, Website Poll, Chatbot and Admin Operations",
        "First-party visitor support, feedback reporting, response-path optimization, and application-layer status",
    )
    add_heading(doc, "Executive milestone")
    add_body(doc, "Toolkit gained a first-party support workflow across the redesigned website. Visitors can request assistance and rate the website; authorized administrators can manage enquiries, poll results, chatbot content, rollout state, and analytics from Toolkit Control.")
    add_heading(doc, "Completed modules")
    add_bullets(doc, [
        "Private WordPress enquiry records with visitor contact, source page, message, date, and new/in-progress/resolved workflow.",
        "Configurable five-point website poll covering design, navigation, content, speed, and mobile improvements.",
        "Poll reporting for response totals, average rating, rating distribution, and improvement-area mentions.",
        "Administrator controls for assistant availability, welcome text, common answers, poll availability, poll question, and guidance.",
        "Operational Toolkit Control overview with enquiry count, poll count, assistant state, rollout controls, system status, and dedicated analytics link.",
        "Assistant made available across redesigned pages instead of only the homepage.",
    ])
    add_heading(doc, "Performance improvements")
    add_bullets(doc, [
        "Embedded chatbot configuration into the page payload, removing the first-open REST configuration delay.",
        "Moved enquiry email notification to a scheduled WordPress event after record storage so SMTP does not block the visitor response.",
        "Moved 30-day metric aggregation away from the first Toolkit Control screen into the dedicated Site analytics page.",
        "Measured the optimized labelled demo enquiry at HTTP 201 in 1.44 seconds.",
    ])
    add_heading(doc, "Security and privacy controls")
    add_control_table(doc, [
        ("Public record exposure", "Disabled; enquiry and poll records are private"),
        ("Request validation", "Same-site Origin/Referer enforcement"),
        ("Abuse controls", "Per-IP rate limits, honeypot, and field limits"),
        ("Input handling", "Allow-lists and WordPress sanitization"),
        ("Enquiry consent", "Required before storage"),
        ("Administrator access", "manage_options plus WordPress nonces"),
        ("Visitor profiling", "No raw IP, replay, fingerprint, or keystroke capture"),
    ])
    add_heading(doc, "Validation record")
    add_bullets(doc, [
        "PHP and JavaScript syntax checks passed.",
        "Demo enquiry persistence returned HTTP 201.",
        "Demo poll persistence returned HTTP 201 after correcting the internal post-type key to WordPress's 20-character limit.",
        "Invalid-origin submissions returned HTTP 403 on demo and production.",
        "Production browser interaction passed panel opening, six actions, complete enquiry fields, five ratings, five improvement areas, and poll comment.",
        "Production desktop/mobile rendering, support configuration, admin stylesheet, security headers, and LiteSpeed cache response passed.",
        "Three clearly labelled system-test records remain on demo for administrator inspection: two enquiries and one poll response.",
    ])
    add_evidence(
        doc,
        ROOT / "rollbacks" / "production-20260723-support" / "screenshots" / "support-home-desktop.png",
        "Main-domain desktop verification with the Toolkit support launcher, captured 23 July 2026.",
    )
    add_evidence(
        doc,
        ROOT / "rollbacks" / "production-20260723-support" / "screenshots" / "support-home-mobile.png",
        "Main-domain mobile verification with the compact Toolkit support launcher, captured 23 July 2026.",
    )
    add_heading(doc, "Application-layer status - not deployed")
    add_bullets(doc, [
        "Local implementation includes a branded six-step form and same-origin Mzizi campus, course, intake, and submission adapters.",
        "Implemented safeguards include field allow-listing, server-side selection revalidation, rate limiting, honeypot, Turnstile hooks, and an explicit submission feature flag.",
        "Current fail-closed behavior opens the official Mzizi portal and confirms that locally entered data was not transmitted.",
        "Production activation still requires approved Mzizi authorization, environment-managed Turnstile keys, non-production contract testing, accessibility and failure-path testing, privacy approval, and rollback acceptance.",
    ])
    add_heading(doc, "Next dated workstream")
    add_bullets(doc, [
        "Privacy-conscious user journey and conversion reporting with defined retention and role access.",
        "Google Search Console and Bing Webmaster integration, indexing coverage, branded/non-branded query reporting, and technical SEO monitoring.",
        "Authenticated administrator visual acceptance using an existing WordPress administrator account.",
    ])
    doc.save(OUTPUT / "Toolkit_Weekly_Milestone_12_Support_Admin_and_Application_Status_23_July_2026.docx")


if __name__ == "__main__":
    OUTPUT.mkdir(parents=True, exist_ok=True)
    production_report()
    support_report()
    print("Generated WMR-11 and WMR-12")

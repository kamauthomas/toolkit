#!/usr/bin/env python3
from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/Toolkit_Strategic_ICT_Roadmap_2026_2027.md"
OUT = ROOT / "reports/Toolkit_Strategic_ICT_Roadmap_2026_2027.docx"
LOGO = ROOT / "wp-content/themes/eduma-child/assets/images/toolkit-logo.png"
TEAL, OLIVE, ORANGE, DARK = "006A68", "969E2A", "FF6600", "262B2A"


def rgb(value):
    return RGBColor.from_string(value)


def shade(cell, value):
    props = cell._tc.get_or_add_tcPr()
    item = OxmlElement("w:shd")
    item.set(qn("w:fill"), value)
    props.append(item)


def clean(text):
    return re.sub(r"\*\*(.*?)\*\*", r"\1", text).replace("`", "")


doc = Document()
section = doc.sections[0]
section.top_margin = section.bottom_margin = Inches(0.62)
section.left_margin = section.right_margin = Inches(0.68)
normal = doc.styles["Normal"]
normal.font.name = "Aptos"
normal.font.size = Pt(9.2)
normal.font.color.rgb = rgb(DARK)
for name, size, color in (
    ("Title", 24, TEAL),
    ("Heading 1", 16, TEAL),
    ("Heading 2", 13, OLIVE),
    ("Heading 3", 11, ORANGE),
):
    style = doc.styles[name]
    style.font.name = "Aptos Display"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = rgb(color)

header = section.header.paragraphs[0]
header.text = "TOOLKIT AFRICA  /  STRATEGIC ICT ROADMAP  /  2026–2027"
header.runs[0].font.size = Pt(8)
header.runs[0].font.bold = True
header.runs[0].font.color.rgb = rgb(TEAL)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("Toolkit Africa | Draft for Management Review | 28 July 2026").font.size = Pt(7.5)

cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.add_run().add_picture(str(LOGO), width=Inches(1.55))
title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("Strategic ICT Roadmap and Implementation Plan")
subtitle = doc.add_paragraph("AUGUST 2026 – JULY 2027")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].bold = True
subtitle.runs[0].font.color.rgb = rgb(ORANGE)
focus = doc.add_paragraph(
    "Eleven strategic initiatives | 90-day mobilisation | 12-month delivery | Costs and resources"
)
focus.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

lines = SOURCE.read_text().splitlines()
i = 1
while i < len(lines):
    line = lines[i].strip()
    if not line:
        i += 1
        continue
    if line.startswith("|") and i + 1 < len(lines) and set(lines[i + 1].replace("|", "").replace("-", "").replace(":", "").strip()) == set():
        headers = [clean(x.strip()) for x in line.strip("|").split("|")]
        i += 2
        rows = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            rows.append([clean(x.strip()) for x in lines[i].strip().strip("|").split("|")])
            i += 1
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for cell, value in zip(table.rows[0].cells, headers):
            shade(cell, TEAL)
            run = cell.paragraphs[0].add_run(value)
            run.bold = True
            run.font.color.rgb = rgb("FFFFFF")
        for values in rows:
            cells = table.add_row().cells
            for cell, value in zip(cells, values):
                cell.text = value
        continue
    if line.startswith("### "):
        doc.add_heading(clean(line[4:]), level=3)
    elif line.startswith("## "):
        doc.add_heading(clean(line[3:]), level=1)
    elif line.startswith("# "):
        pass
    elif re.match(r"^\d+\.\s", line):
        doc.add_paragraph(clean(re.sub(r"^\d+\.\s+", "", line)), style="List Number")
    elif line.startswith("- "):
        doc.add_paragraph(clean(line[2:]), style="List Bullet")
    else:
        paragraph = line
        while i + 1 < len(lines):
            nxt = lines[i + 1].strip()
            if not nxt or nxt.startswith(("#", "-", "|")) or re.match(r"^\d+\.\s", nxt):
                break
            paragraph += " " + nxt
            i += 1
        doc.add_paragraph(clean(paragraph))
    i += 1

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)

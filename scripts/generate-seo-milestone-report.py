#!/usr/bin/env python3
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "reports/weekly-milestones/Toolkit_Weekly_Milestone_08_SEO_Milestones_and_Fixes_17_July_2026.docx"
LOGO = ROOT / "wp-content/themes/eduma-child/assets/images/toolkit-logo.png"
BASELINE = Path("/tmp/codex-clipboard-no3vwc.png")
AFTER = Path("/tmp/codex-clipboard-ei7yeM.png")
DETAIL = Path("/tmp/codex-clipboard-vmdhkD.png")
OLIVE, ORANGE, TEAL, DARK, LIGHT = "969E2A", "FF6600", "006A68", "262B2A", "F4F6F2"


def color(value):
    return RGBColor.from_string(value)


def shade(cell, value):
    props = cell._tc.get_or_add_tcPr()
    item = OxmlElement("w:shd")
    item.set(qn("w:fill"), value)
    props.append(item)


def heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def bullets(doc, values):
    for value in values:
        doc.add_paragraph(value, style="List Bullet")


def metric_table(doc):
    rows = [
        ("Mobile performance", "58", "84", "+26 points"),
        ("First Contentful Paint", "2.0 s", "1.7 s", "15% faster"),
        ("Largest Contentful Paint", "15.5 s", "3.1 s", "80% faster"),
        ("Total Blocking Time", "350 ms", "60 ms", "83% lower"),
        ("Cumulative Layout Shift", "0.007", "0.045", "Still within pass threshold"),
        ("Best practices", "92", "96", "+4 points"),
        ("SEO (PageSpeed)", "92", "92", "Maintained"),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for cell, label in zip(table.rows[0].cells, ("Metric", "Baseline", "After fixes", "Outcome")):
        shade(cell, TEAL)
        run = cell.paragraphs[0].add_run(label)
        run.bold = True
        run.font.color.rgb = color("FFFFFF")
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = value
        cells[3].paragraphs[0].runs[0].font.color.rgb = color(OLIVE)


def add_image(doc, path, caption):
    if not path.exists():
        doc.add_paragraph(f"Evidence unavailable: {path.name}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(6.8))
    c = doc.add_paragraph(caption)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.runs[0].italic = True
    c.runs[0].font.size = Pt(8)


doc = Document()
section = doc.sections[0]
section.top_margin = section.bottom_margin = Inches(0.65)
section.left_margin = section.right_margin = Inches(0.72)
normal = doc.styles["Normal"]
normal.font.name = "Aptos"
normal.font.size = Pt(9.5)
normal.font.color.rgb = color(DARK)
for name, size, value in (("Title", 27, TEAL), ("Heading 1", 17, TEAL), ("Heading 2", 12.5, OLIVE)):
    style = doc.styles[name]
    style.font.name = "Aptos Display"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = color(value)

header = section.header.paragraphs[0]
header.text = "TOOLKIT AFRICA  /  WEBSITE MODERNISATION  /  WMR-08"
header.runs[0].bold = True
header.runs[0].font.color.rgb = color(TEAL)
header.runs[0].font.size = Pt(8)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("Toolkit Africa | SEO Milestones and Fixes | 17 July 2026 | Confidential project record").font.size = Pt(7.5)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(str(LOGO), width=Inches(1.5))
p = doc.add_paragraph("WEBSITE MODERNISATION PROGRAMME")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].bold = True
p.runs[0].font.color.rgb = color(ORANGE)
title = doc.add_paragraph("SEO Milestones and Fixes", style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
period = doc.add_paragraph("Weekly Milestone Report 08 | 17 July 2026")
period.alignment = WD_ALIGN_PARAGRAPH.CENTER
period.runs[0].bold = True
period.runs[0].font.color.rgb = color(OLIVE)

control = doc.add_table(rows=4, cols=2)
control.alignment = WD_TABLE_ALIGNMENT.CENTER
for index, values in enumerate((("Report", "WMR-08"), ("Target", "https://demo.toolkitafrica.ac.ke/"), ("Release commits", "68dce70, c78c4d5, 6ef80f3"), ("Status", "Implemented, deployed and live-verified"))):
    control.cell(index, 0).text, control.cell(index, 1).text = values
    shade(control.cell(index, 0), TEAL)
    control.cell(index, 0).paragraphs[0].runs[0].font.color.rgb = color("FFFFFF")
    control.cell(index, 0).paragraphs[0].runs[0].bold = True

doc.add_paragraph()
box = doc.add_table(rows=1, cols=1)
shade(box.cell(0, 0), LIGHT)
box.cell(0, 0).text = "Weekly outcome\nThe mobile PageSpeed score improved from 58 to 84. LCP reduced from 15.5 seconds to 3.1 seconds through responsive WebP delivery, legacy asset removal, targeted preloading, click-to-load video facades, corrected metadata, and verified LiteSpeed caching."
box.cell(0, 0).paragraphs[0].runs[0].font.color.rgb = color(OLIVE)
doc.add_page_break()

heading(doc, "1. Executive Summary")
doc.add_paragraph("This milestone addressed mobile performance, technical SEO, media delivery, structured metadata, and response-security findings raised during external audits of the demo environment. Corrections were implemented in the child theme, deployed through guarded FTPS, cache-purged, and verified against live response headers and HTML.")
heading(doc, "Measured Improvement", 2)
metric_table(doc)
heading(doc, "Primary Deliverables", 2)
bullets(doc, [
    "Responsive desktop and mobile WebP hero assets with media-specific high-priority preloads.",
    "Removal of unused Elementor, Contact Form 7, Thim eKit, Sina, and associated runtime assets from the custom homepage.",
    "Click-to-load YouTube facades for the pathways video and all three graduate testimonials.",
    "Keyword-focused title, standard meta description, self-canonical URL, Open Graph description, and dedicated 1200 x 630 social image.",
    "Meaningful image alternative text, HSTS response header, and live validation of SPF, robots.txt, sitemap, analytics, and cache behaviour.",
])
doc.add_page_break()

heading(doc, "2. Performance Evidence")
add_image(doc, BASELINE, "Baseline mobile PageSpeed result: performance 58 and LCP 15.5 seconds.")
doc.add_page_break()
add_image(doc, AFTER, "Post-deployment mobile PageSpeed result: performance 84 and LCP 3.1 seconds.")
add_image(doc, DETAIL, "Post-deployment diagnostics showing reduced blocking and remaining font/render-path opportunities.")
doc.add_page_break()

heading(doc, "3. SEO and Technical Corrections")
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
for cell, label in zip(table.rows[0].cells, ("Finding", "Action", "Verified state")):
    shade(cell, TEAL)
    cell.text = label
    cell.paragraphs[0].runs[0].font.color.rgb = color("FFFFFF")
for row in (
    ("Missing meta description", "Added curated homepage description through Yoast filters.", "Present in live HTML."),
    ("Weak title/keyword alignment", "Replaced generic title with practical-skills and Kenya context.", "Live title updated."),
    ("Canonical mismatch", "Canonical and Open Graph URL now use the active environment home URL.", "Demo self-canonical verified."),
    ("Image formats and sizing", "Introduced mobile/desktop WebP assets and a dedicated social image.", "WebP delivery verified."),
    ("Missing alt attributes", "Added descriptive alt text to video and story images.", "No empty homepage alt attributes."),
    ("Heavy third-party embeds", "Replaced initial iframes with accessible click-to-load facades.", "Zero initial YouTube iframes."),
    ("Missing HSTS", "Added one-year includeSubDomains policy on secure frontend responses.", "Header verified live."),
):
    cells = table.add_row().cells
    for cell, value in zip(cells, row):
        cell.text = value

heading(doc, "Audit Claims Rechecked", 2)
bullets(doc, [
    "SPF was reported missing but DNS returns v=spf1 include:spf.protection.outlook.com -all.",
    "robots.txt was reported as blocking content but its public user-agent group has an empty Disallow directive.",
    "The sitemap index returns HTTP 200.",
    "Unsafe target=_blank links were not present in current homepage HTML.",
    "Analytics is implemented through Toolkit's first-party metrics module rather than Google Analytics.",
])
doc.add_page_break()

heading(doc, "4. Deployment, Controls and Remaining Work")
heading(doc, "Deployment Controls", 2)
bullets(doc, [
    "Only changed child-theme files were uploaded; WordPress core, plugins, database, and uploads were untouched.",
    "Persistent rollback copies were refreshed under rollbacks/latest-demo before each release.",
    "Temporary token-protected cache purge files were removed immediately and verified as HTTP 404.",
    "LiteSpeed behaviour was verified as MISS on the first request and HIT on the next request.",
])
heading(doc, "Remaining Opportunities", 2)
bullets(doc, [
    "Reduce remaining font-display delay by consolidating parent-theme icon/font families.",
    "Further isolate critical header/footer CSS from the parent Eduma stylesheet after full regression testing.",
    "Retest PageSpeed across multiple runs because lab scores vary with network and test location.",
    "Review accessibility findings individually; the current PageSpeed accessibility score remains 80.",
    "Ads.txt is only required if Toolkit becomes an advertising publisher; it is not an SEO requirement for this site.",
    "The visible office email is retained intentionally for contact usability; anti-spam controls should not hide essential contact information.",
])
heading(doc, "Acceptance Statement", 2)
doc.add_paragraph("The high-impact performance and technical SEO corrections in this milestone are deployed and verified on the demo environment. The measured improvement is material and the remaining work is primarily parent-theme font/CSS reduction, accessibility refinement, and repeated monitoring before main-domain cutover.")

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)

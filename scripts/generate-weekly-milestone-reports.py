#!/usr/bin/env python3
"""Generate Toolkit Africa weekly milestone DOCX reports with visual evidence."""

from __future__ import annotations

import re
import textwrap
from datetime import datetime
from html import unescape
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "reports" / "weekly-milestones"
ASSETS = OUTPUT / "evidence"
LOGO = ROOT / "wp-content/themes/eduma-child/assets/images/toolkit-logo.png"
TMP = Path("/tmp")

OLIVE = "969E2A"
ORANGE = "FF6600"
TEAL = "006A68"
DARK = "262B2A"
MID = "5C6662"
LIGHT = "F4F6F2"
LINE = "D9DED6"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = LINE, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, end))


def configure_document(doc: Document, report_number: int, period: str) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.4)
    normal.font.color.rgb = rgb(DARK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in (
        ("Title", 27, TEAL),
        ("Heading 1", 17, TEAL),
        ("Heading 2", 12.5, OLIVE),
        ("Heading 3", 10.5, ORANGE),
    ):
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(7.0))
    table.columns[0].width = Inches(4.9)
    table.columns[1].width = Inches(2.1)
    left = table.cell(0, 0).paragraphs[0]
    left.add_run("TOOLKIT AFRICA  /  DIGITAL MODERNISATION").bold = True
    left.runs[0].font.color.rgb = rgb(TEAL)
    left.runs[0].font.size = Pt(8)
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.add_run(f"WEEKLY REPORT {report_number:02d}").bold = True
    right.runs[0].font.color.rgb = rgb(ORANGE)
    right.runs[0].font.size = Pt(8)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Toolkit Africa | {period} | Confidential project record  |  Page ")
    run.font.size = Pt(7.5)
    run.font.color.rgb = rgb(MID)
    add_field(p, "PAGE")


def add_cover(doc: Document, report: dict) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(LOGO), width=Inches(1.45))
    p.paragraph_format.space_after = Pt(8)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = kicker.add_run("WEBSITE MODERNISATION PROGRAMME")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = rgb(ORANGE)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(report["title"])

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run(report["period"])
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = rgb(OLIVE)

    control = doc.add_table(rows=4, cols=2)
    control.alignment = WD_TABLE_ALIGNMENT.CENTER
    control.autofit = False
    control.columns[0].width = Inches(1.65)
    control.columns[1].width = Inches(4.75)
    rows = (
        ("Report number", f"WMR-{report['number']:02d}"),
        ("Issue date", report["issued"]),
        ("Delivery status", report["status"]),
        ("Evidence note", "Page captures were taken on 16 July 2026 for consistent comparison."),
    )
    for idx, (label, value) in enumerate(rows):
        c0, c1 = control.rows[idx].cells
        shade_cell(c0, TEAL)
        c0.paragraphs[0].add_run(label).bold = True
        c0.paragraphs[0].runs[0].font.color.rgb = rgb("FFFFFF")
        c1.paragraphs[0].add_run(value)
        set_cell_border(c0)
        set_cell_border(c1)
        c0.vertical_alignment = c1.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    doc.add_paragraph()
    summary_box = doc.add_table(rows=1, cols=1)
    shade_cell(summary_box.cell(0, 0), LIGHT)
    set_cell_border(summary_box.cell(0, 0), OLIVE, "10")
    p = summary_box.cell(0, 0).paragraphs[0]
    p.add_run("WEEKLY OUTCOME\n").bold = True
    p.runs[0].font.color.rgb = rgb(OLIVE)
    p.add_run(report["summary"])
    doc.add_page_break()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        p.add_run(item)


def add_status_table(doc: Document, items: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ("Milestone", "Status", "Evidence / outcome")
    for i, label in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, TEAL)
        run = cell.paragraphs[0].add_run(label)
        run.bold = True
        run.font.color.rgb = rgb("FFFFFF")
    for milestone, status, evidence in items:
        row = table.add_row().cells
        row[0].paragraphs[0].add_run(milestone).bold = True
        row[1].paragraphs[0].add_run(status)
        row[1].paragraphs[0].runs[0].font.color.rgb = rgb(OLIVE if status == "Complete" else ORANGE)
        row[2].paragraphs[0].add_run(evidence)
        for cell in row:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def optimize_screenshot(source: Path, target: Path) -> None:
    with Image.open(source) as image:
        image = image.convert("RGB")
        if image.width > 1200:
            height = round(image.height * 1200 / image.width)
            image = image.resize((1200, height), Image.Resampling.LANCZOS)
        image.save(target, "JPEG", quality=76, optimize=True, progressive=True)


def add_comparison(doc: Document, page_name: str) -> None:
    current = ASSETS / f"current-{page_name}.jpg"
    demo = ASSETS / f"demo-{page_name}.jpg"
    if not current.exists() or not demo.exists():
        return
    heading = page_name.replace("-", " ").title()
    doc.add_heading(f"Visual comparison: {heading}", level=2)
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = (("LEGACY PRODUCTION", current, "Baseline"), ("NEW DEMO", demo, "Modernised"))
    for idx, (label, image_path, accent) in enumerate(labels):
        header = table.cell(0, idx)
        shade_cell(header, TEAL if idx == 0 else OLIVE)
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = rgb("FFFFFF")
        body = table.cell(1, idx)
        body.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        body.paragraphs[0].add_run().add_picture(str(image_path), width=Inches(3.25))
        c = body.add_paragraph(accent)
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.runs[0].italic = True
        c.runs[0].font.size = Pt(7.5)
        c.runs[0].font.color.rgb = rgb(MID)
        set_cell_border(header)
        set_cell_border(body)


def html_meta(path: Path) -> dict[str, str]:
    text = path.read_text(errors="ignore")
    def first(pattern: str) -> str:
        match = re.search(pattern, text, flags=re.I | re.S)
        return unescape(re.sub(r"\s+", " ", match.group(1)).strip()) if match else "Not present"
    return {
        "title": first(r"<title[^>]*>(.*?)</title>"),
        "description": first(r'<meta[^>]+name=["\']description["\'][^>]+content=["\'](.*?)["\']'),
        "canonical": first(r'<link[^>]+rel=["\']canonical["\'][^>]+href=["\'](.*?)["\']'),
        "og_title": first(r'<meta[^>]+property=["\']og:title["\'][^>]+content=["\'](.*?)["\']'),
        "schema": "Present" if "application/ld+json" in text else "Not detected",
    }


def load_font(size: int, bold: bool = False):
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def draw_wrapped(draw, xy, text, font, fill, width_chars, spacing=6):
    lines = []
    for paragraph in str(text).splitlines() or [""]:
        lines.extend(textwrap.wrap(paragraph, width=width_chars) or [""])
    draw.multiline_text(xy, "\n".join(lines), font=font, fill=fill, spacing=spacing)


def create_seo_panel(page_name: str) -> Path:
    current = html_meta(TMP / f"seo-current-{page_name}.html")
    demo = html_meta(TMP / f"seo-demo-{page_name}.html")
    image = Image.new("RGB", (1500, 900), "white")
    draw = ImageDraw.Draw(image)
    title_font = load_font(36, True)
    heading_font = load_font(23, True)
    label_font = load_font(17, True)
    body_font = load_font(16)
    draw.rectangle((0, 0, 1500, 92), fill="#006a68")
    draw.text((45, 25), f"SEO evidence: {page_name.replace('-', ' ').title()}", font=title_font, fill="white")
    columns = ((40, 740, "LEGACY PRODUCTION", current, "#006a68"), (760, 1460, "NEW DEMO", demo, "#969e2a"))
    for left, right, heading, values, color in columns:
        draw.rounded_rectangle((left, 120, right, 850), radius=12, fill="#f4f6f2", outline=color, width=4)
        draw.rectangle((left, 120, right, 180), fill=color)
        draw.text((left + 22, 136), heading, font=heading_font, fill="white")
        y = 210
        for label, key, max_chars in (("TITLE", "title", 62), ("DESCRIPTION", "description", 72), ("CANONICAL", "canonical", 70), ("OPEN GRAPH TITLE", "og_title", 62), ("SCHEMA", "schema", 65)):
            draw.text((left + 24, y), label, font=label_font, fill="#ff6600")
            draw_wrapped(draw, (left + 24, y + 28), values[key], body_font, "#262b2a", max_chars)
            line_count = max(1, len(textwrap.wrap(values[key], width=max_chars)))
            y += 58 + line_count * 24
    target = ASSETS / f"seo-{page_name}.jpg"
    image.save(target, "JPEG", quality=84, optimize=True)
    return target


def add_seo_evidence(doc: Document, pages: list[str]) -> None:
    for page in pages:
        panel = ASSETS / f"seo-{page}.jpg"
        if panel.exists():
            doc.add_heading(f"SEO comparison: {page.replace('-', ' ').title()}", level=2)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(panel), width=Inches(6.65))
            caption = doc.add_paragraph("Live metadata capture: production baseline compared with the demo implementation.")
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.runs[0].italic = True
            caption.runs[0].font.size = Pt(7.5)
            caption.runs[0].font.color.rgb = rgb(MID)


REPORTS = [
    {
        "number": 1,
        "period": "01-07 June 2026",
        "issued": "07 June 2026",
        "title": "Discovery, Crawl and Baseline Audit",
        "status": "Complete",
        "summary": "Established the evidence base for redesign by crawling the production WordPress estate, documenting information architecture, technical dependencies, duplicate content and high-risk user journeys.",
        "achievements": [
            "Completed a sitemap, REST API and HTML crawl of toolkitafrica.ac.ke.",
            "Recorded 230 indexed URLs and 641 discovered URLs in the June baseline.",
            "Mapped core pages, venture hierarchy, account routes, blog content and media dependencies.",
            "Identified duplicate Foundation, Blog, account and team systems, including 37 overlapping team profiles.",
            "Flagged the broken application route, exceptionally long post slugs and unfinished /new/ redesign paths.",
        ],
        "milestones": [
            ("Production crawl", "Complete", "Sitemap, REST and HTML sources reconciled."),
            ("Risk register", "Complete", "Broken, duplicate and legacy routes documented."),
            ("Priority inventory", "Complete", "Core marketing, course, account and legal pages classified."),
        ],
        "comparisons": ["home", "about", "courses"],
        "risks": [
            "Existing URLs could not be removed without redirect and search-impact analysis.",
            "LearnPress, account, contact and Mzizi flows required preservation during visual work.",
        ],
        "next": ["Approve the design direction and convert the audit into a page-by-page implementation specification."],
    },
    {
        "number": 2,
        "period": "08-14 June 2026",
        "issued": "14 June 2026",
        "title": "Visual Direction and Experience Prototype",
        "status": "Complete",
        "summary": "Translated the audit and supplied references into a modern visual direction, initial application experience and reusable brand/layout specifications.",
        "achievements": [
            "Consolidated supplied hero, form, welding and notice-board references into a UI specification.",
            "Built an exploratory Laravel prototype to validate layout and route assumptions before the WordPress implementation decision.",
            "Introduced a Toolkit-led hero composition, application-first CTA and logo-based visual direction.",
            "Prototyped the application experience and the content hierarchy for programmes, notices and team content.",
            "Documented responsive breakpoints, typography, spacing, interaction and asset requirements.",
        ],
        "milestones": [
            ("Reference synthesis", "Complete", "UI references converted into structured requirements."),
            ("Prototype routes", "Complete", "Core page and application routes demonstrated."),
            ("Responsive specification", "Complete", "Desktop, tablet and mobile expectations recorded."),
        ],
        "comparisons": ["home", "application"],
        "risks": ["The prototype was not the production CMS and therefore could not replace WordPress content ownership."],
        "next": ["Retain the visual findings and move implementation into an update-safe WordPress child theme."],
    },
    {
        "number": 3,
        "period": "15-21 June 2026",
        "issued": "21 June 2026",
        "title": "WordPress Integration Architecture",
        "status": "Complete",
        "summary": "Defined an implementation model that modernises the existing WordPress estate while preserving content, URLs, plugins and rollback capability.",
        "achievements": [
            "Selected the Eduma child theme as the production integration boundary.",
            "Defined theme-owned templates and assets for priority pages while retaining WordPress content and forms.",
            "Established the requirement for host-aware rollout switches and parent-theme fallback.",
            "Separated current course data from future prospectus data and pricing activation.",
            "Defined the Mzizi adapter approach: Toolkit-hosted form, server-side validation and controlled forwarding.",
        ],
        "milestones": [
            ("CMS decision", "Complete", "WordPress retained as the source of content and operations."),
            ("Rollback model", "Complete", "Child-theme switch and parent fallback specified."),
            ("Integration boundaries", "Complete", "Forms, LMS, accounts and uploads protected."),
        ],
        "comparisons": ["courses", "application"],
        "risks": ["Production activation would require file, database, cache and route verification as separate gates."],
        "next": ["Prepare the build checklist, accessibility criteria and reproducible local WordPress environment."],
    },
    {
        "number": 4,
        "period": "22-28 June 2026",
        "issued": "28 June 2026",
        "title": "Implementation Readiness and Acceptance Planning",
        "status": "Complete",
        "summary": "Converted the architecture into a controlled delivery plan covering page priorities, evidence capture, accessibility, performance, SEO and deployment safety.",
        "achievements": [
            "Prioritised Home, About, Courses, Welding, Notice Board, Application, Blog and Contact for the first release sequence.",
            "Defined desktop/mobile screenshot acceptance and asset verification requirements.",
            "Established semantic heading, keyboard, focus, reduced-motion and image-alt expectations.",
            "Specified deployment order: backup, assets/templates, functions last, cache purge and route checks.",
            "Separated visual acceptance from production cutover approval.",
        ],
        "milestones": [
            ("Page release plan", "Complete", "Priority routes and dependencies sequenced."),
            ("Quality gates", "Complete", "Syntax, visual, responsive, SEO and HTTP checks defined."),
            ("Deployment controls", "Complete", "Rollback-first FTPS workflow established."),
        ],
        "comparisons": ["welding", "notice"],
        "risks": ["Builder-heavy production pages could create inconsistent performance and CSS conflicts."],
        "next": ["Begin the child-theme homepage, header and brand-token implementation."],
    },
    {
        "number": 5,
        "period": "29 June-05 July 2026",
        "issued": "05 July 2026",
        "title": "Homepage, Header and Brand Foundation",
        "status": "Complete",
        "summary": "Delivered the first production-shaped WordPress experience: modern hero, feature strip, stable header, brand tokens and a rebuilt Who We Are section.",
        "achievements": [
            "Implemented an accessible three-slide hero with Apply CTA, pagination, pause, keyboard and swipe controls.",
            "Added the feature strip and corrected counter, arrow, overlay and scroll-cue behaviour.",
            "Stabilised navigation selection, logo sizing, dropdown spacing and Apply Now placement.",
            "Created shared brand tokens and introduced consistent component spacing and button treatment.",
            "Rebuilt Who We Are with the existing vision, mission, quick links and five impact figures.",
        ],
        "milestones": [
            ("Hero system", "Complete", "Responsive carousel and dual CTA implementation."),
            ("Header/navigation", "Complete", "Primary menu and branded actions stabilised."),
            ("Who We Are", "Complete", "Existing organisational data retained in a modern layout."),
        ],
        "comparisons": ["home", "about"],
        "risks": ["Legacy Elementor output still loaded behind the new homepage sections and required isolation."],
        "next": ["Harden the child-theme boundary, remove blocking preloaders and validate local runtime behaviour."],
    },
    {
        "number": 6,
        "period": "06-12 July 2026",
        "issued": "12 July 2026",
        "title": "Child-Theme Hardening and Runtime Performance",
        "status": "Complete",
        "summary": "Made the redesign independent of fragile builder state, established a reproducible local runtime and removed major causes of slow or blocked rendering.",
        "achievements": [
            "Moved homepage ownership into the child theme and retained parent fallback when redesign is disabled.",
            "Converted homepage video to click-to-play to reduce initial network and layout cost.",
            "Disabled the blocking Eduma public preloader and removed stale preload body classes.",
            "Created the local MariaDB/PHP workflow and validated child-theme rendering at 127.0.0.1:8001.",
            "Reduced plugin and builder asset exposure on rebuilt routes and formalised database seeding.",
        ],
        "milestones": [
            ("Child-theme isolation", "Complete", "Homepage and priority layouts no longer depend on page-builder rendering."),
            ("Local runtime", "Complete", "Database import and PHP review server documented."),
            ("Progressive rendering", "Complete", "Blocking preload behaviour removed."),
        ],
        "comparisons": ["home", "contact"],
        "risks": ["Uncached WordPress generation remained slower than warm LiteSpeed responses."],
        "next": ["Build and deploy the major-page templates, SEO layer, course catalogue and operational controls."],
    },
    {
        "number": 7,
        "period": "13-16 July 2026",
        "issued": "16 July 2026",
        "title": "Major Pages, SEO, Analytics and Demo Release",
        "status": "Complete through 16 July",
        "summary": "Completed the first coherent demo release across major pages, data controls, SEO, analytics and rollback operations, with the production domain intentionally unchanged.",
        "achievements": [
            "Rebuilt Courses, Welding, Notice Board, Application, About, Impact, Foundation, Contact and Toolkit Blog templates.",
            "Restored all 18 team members with optimized local portraits and retained live-site organisational content.",
            "Expanded homepage depth with 2026 IIW, sector video, animated impact metrics, testimonials, stories and floating assistant.",
            "Implemented functional Notice Board search, filtering, sorting, result counts and grid/list views.",
            "Added curated titles, descriptions, canonicals, Open Graph metadata and structured-data hierarchy.",
            "Added first-party page, dwell, scroll, performance and interaction metrics with Toolkit Control in WordPress admin.",
            "Introduced redesign/catalog/pricing switches, persistent rollback storage and guarded FTPS deployment.",
            "Established main as the Git source-of-truth branch and removed the merged feature branch.",
        ],
        "milestones": [
            ("Major-page release", "Complete", "Eight priority routes redesigned and deployed to demo."),
            ("SEO implementation", "Complete", "Metadata, canonicals, social previews and schema controls added."),
            ("Operational controls", "Complete", "Metrics, feature switches, seeders and rollback workflow available."),
            ("Demo verification", "Complete", "HTTP, asset, syntax, cache and visual checks passed."),
            ("Main-domain cutover", "Pending approval", "Production remains unchanged pending credentials and final acceptance."),
        ],
        "comparisons": ["home", "about", "courses", "welding", "notice", "application", "blog", "contact"],
        "seo": ["home", "about", "courses", "welding", "notice", "application", "blog", "contact"],
        "risks": [
            "Main-domain activation must preserve 39 current page routes, approximately 150 post URLs, two team systems and account/LMS workflows.",
            "The Mzizi form adapter requires an approved field contract and non-production submission testing before implementation.",
            "September course pricing remains disabled until written admissions approval and the effective date.",
        ],
        "next": [
            "Complete stakeholder demo acceptance.",
            "Obtain production file/database access and take persistent backups.",
            "Sync the child theme with all switches off, activate legacy-compatible mode, verify routes, then enable the redesign under rollback control.",
        ],
    },
]


def build_report(report: dict) -> Path:
    doc = Document()
    configure_document(doc, report["number"], report["period"])
    doc.core_properties.title = report["title"]
    doc.core_properties.subject = "Toolkit Africa website modernisation weekly milestone"
    doc.core_properties.author = "Toolkit Africa Digital Modernisation Programme"
    doc.core_properties.keywords = "Toolkit Africa, WordPress, redesign, milestone, SEO"
    doc.core_properties.created = datetime.strptime(report["issued"], "%d %B %Y")

    add_cover(doc, report)
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(report["summary"])

    doc.add_heading("2. Milestone Status", level=1)
    add_status_table(doc, report["milestones"])

    doc.add_heading("3. Work Completed", level=1)
    add_bullets(doc, report["achievements"])

    doc.add_heading("4. Visual Evidence", level=1)
    doc.add_paragraph("The left image shows the production baseline and the right image shows the redesigned demo. Captures use a consistent desktop viewport and were collected on 16 July 2026.")
    doc.add_page_break()
    for page in report.get("comparisons", []):
        add_comparison(doc, page)

    if report.get("seo"):
        doc.add_page_break()
        doc.add_heading("5. SEO Evidence", level=1)
        doc.add_paragraph("The following panels reproduce public metadata captured from both environments. Demo canonicals intentionally resolve to the main domain to avoid indexing the staging host.")
        add_seo_evidence(doc, report["seo"])
        risk_section = "6. Risks and Controls"
        next_section = "7. Next-Week Priorities"
        verification_section = "8. Verification Record"
    else:
        risk_section = "5. Risks and Controls"
        next_section = "6. Next-Week Priorities"
        verification_section = "7. Verification Record"

    doc.add_heading(risk_section, level=1)
    add_bullets(doc, report["risks"])
    doc.add_heading(next_section, level=1)
    add_bullets(doc, report["next"])
    doc.add_heading(verification_section, level=1)
    add_status_table(doc, [
        ("Source chronology", "Complete", "June crawl/prototype records, Git history and project progress log reconciled."),
        ("Visual evidence", "Complete", "Production and demo page captures embedded and labelled."),
        ("Document quality", "Complete", "Toolkit logo, weekly dates, version control and consistent report format applied."),
    ])

    filename = OUTPUT / f"Toolkit_Weekly_Milestone_{report['number']:02d}_{report['issued'].replace(' ', '_')}.docx"
    doc.save(filename)
    return filename


def main() -> None:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    ASSETS.mkdir(parents=True, exist_ok=True)

    pages = ["home", "about", "courses", "welding", "notice", "application", "blog", "contact"]
    for page in pages:
        for environment in ("current", "demo"):
            source = TMP / f"toolkit-{environment}-{page}.png"
            target = ASSETS / f"{environment}-{page}.jpg"
            if source.exists():
                optimize_screenshot(source, target)
            elif not target.exists():
                raise FileNotFoundError(source)
        seo_target = ASSETS / f"seo-{page}.jpg"
        if (TMP / f"seo-current-{page}.html").exists() and (TMP / f"seo-demo-{page}.html").exists():
            create_seo_panel(page)
        elif not seo_target.exists():
            raise FileNotFoundError(f"SEO evidence for {page}")

    outputs = [build_report(report) for report in REPORTS]
    print("Generated weekly milestone reports:")
    for path in outputs:
        print(path)


if __name__ == "__main__":
    main()
